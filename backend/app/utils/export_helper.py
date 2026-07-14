import io
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, 
    Paragraph, 
    Spacer, 
    Table, 
    TableStyle, 
    KeepTogether
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
import logging

logger = logging.getLogger(__name__)

def generate_pdf_report(title: str, review_text: str, comparison_table: list, gaps: str, novelty: str) -> io.BytesIO:
    """
    Builds a beautifully styled PDF report from literature review data.
    Returns bytes stream.
    """
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer, 
        pagesize=letter,
        rightMargin=54, 
        leftMargin=54, 
        topMargin=54, 
        bottomMargin=54
    )
    
    styles = getSampleStyleSheet()
    
    # Custom color palette matching our brand (Violet / Slate)
    primary_color = colors.HexColor('#6d28d9')   # Violet-700
    text_color = colors.HexColor('#1e293b')      # Slate-800
    bg_light = colors.HexColor('#f8fafc')        # Slate-50
    
    # Custom styles
    title_style = ParagraphStyle(
        'DocTitle',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=24,
        leading=28,
        textColor=primary_color,
        spaceAfter=15
    )
    
    h2_style = ParagraphStyle(
        'SectionH2',
        parent=styles['Heading2'],
        fontName='Helvetica-Bold',
        fontSize=14,
        leading=18,
        textColor=primary_color,
        spaceBefore=15,
        spaceAfter=8,
        keepWithNext=True
    )
    
    body_style = ParagraphStyle(
        'Body',
        parent=styles['BodyText'],
        fontName='Helvetica',
        fontSize=10.5,
        leading=15,
        textColor=text_color,
        spaceAfter=10
    )
    
    table_text_style = ParagraphStyle(
        'TableText',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=8.5,
        leading=11,
        textColor=text_color
    )
    
    table_header_style = ParagraphStyle(
        'TableHeader',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=8.5,
        leading=11,
        textColor=colors.white
    )
    
    story = []
    
    # 1. Title Page
    story.append(Paragraph(title, title_style))
    story.append(Paragraph("Athena AI - Research Synthesis Platform Report", ParagraphStyle('Sub', parent=body_style, textColor=colors.HexColor('#64748b'), fontName='Helvetica-Oblique')))
    story.append(Spacer(1, 15))
    story.append(Table([[ "" ]], colWidths=[504], rowHeights=[2], style=TableStyle([('BACKGROUND', (0,0), (-1,-1), primary_color)])))
    story.append(Spacer(1, 15))
    
    # 2. Review Text Section
    story.append(Paragraph("Synthesis Review", h2_style))
    story.append(Paragraph(review_text, body_style))
    story.append(Spacer(1, 10))
    
    # 3. Gaps & Novelty
    story.append(Paragraph("Identified Gaps & Proposed Novelty", h2_style))
    story.append(Paragraph(f"<b>Research Gap:</b> {gaps}", body_style))
    story.append(Paragraph(f"<b>Novelty Directives:</b> {novelty}", body_style))
    story.append(Spacer(1, 10))
    
    # 4. Comparison Table (Wrap cell contents in Paragraphs to support auto-wrapping)
    if comparison_table:
        story.append(Paragraph("Comparison Matrix", h2_style))
        
        table_data = [[
            Paragraph("<b>Paper Title</b>", table_header_style),
            Paragraph("<b>Methodology</b>", table_header_style),
            Paragraph("<b>Benchmarks</b>", table_header_style),
            Paragraph("<b>Strengths</b>", table_header_style),
            Paragraph("<b>Weaknesses</b>", table_header_style)
        ]]
        
        for row in comparison_table:
            table_data.append([
                Paragraph(row.get('paper_title', 'Untitled'), table_text_style),
                Paragraph(row.get('method', ''), table_text_style),
                Paragraph(row.get('dataset', ''), table_text_style),
                Paragraph(row.get('strengths', ''), table_text_style),
                Paragraph(row.get('weaknesses', ''), table_text_style)
            ])
            
        col_widths = [110, 100, 94, 100, 100] # total = 504 pt (matching page width minus margins)
        
        t = Table(table_data, colWidths=col_widths, repeatRows=1)
        t.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), primary_color),
            ('ALIGN', (0,0), (-1,-1), 'LEFT'),
            ('VALIGN', (0,0), (-1,-1), 'TOP'),
            ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, bg_light]),
            ('BOX', (0,0), (-1,-1), 0.5, colors.HexColor('#cbd5e1')),
            ('INNERGRID', (0,0), (-1,-1), 0.5, colors.HexColor('#e2e8f0')),
            ('TOPPADDING', (0,0), (-1,-1), 6),
            ('BOTTOMPADDING', (0,0), (-1,-1), 6),
        ]))
        story.append(t)
        
    doc.build(story)
    buffer.seek(0)
    return buffer

def generate_docx_report(title: str, review_text: str, comparison_table: list, gaps: str, novelty: str) -> io.BytesIO:
    """
    Builds a beautifully styled Word Document (.docx) from literature review data.
    Returns bytes stream.
    """
    doc = Document()
    
    # 1. Main title
    doc.add_heading(title, 0)
    doc.add_paragraph("Compiled by Athena AI Research Assistant", style='Subtitle')
    
    # 2. Synthesis Narrative
    doc.add_heading("Comparative Synthesis Review", level=1)
    doc.add_paragraph(review_text)
    
    # 3. Gaps & Novelty
    doc.add_heading("Identified Gaps & Novelty Guidelines", level=1)
    p_gap = doc.add_paragraph()
    p_gap.add_run("Research Gaps: ").bold = True
    p_gap.add_run(gaps)
    
    p_nov = doc.add_paragraph()
    p_nov.add_run("Novelty Directives: ").bold = True
    p_nov.add_run(novelty)
    
    # 4. Table Comparison
    if comparison_table:
        doc.add_heading("Literature Comparison Matrix", level=1)
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Light Shading Accent 1'
        
        # Headers
        hdr_cells = table.rows[0].cells
        headers = ["Paper Title", "Methodology", "Benchmarks", "Strengths", "Weaknesses"]
        for idx, text in enumerate(headers):
            hdr_cells[idx].text = text
            hdr_cells[idx].paragraphs[0].runs[0].font.bold = True
            
        # Rows
        for row in comparison_table:
            row_cells = table.add_row().cells
            row_cells[0].text = row.get('paper_title', 'Untitled')
            row_cells[1].text = row.get('method', '')
            row_cells[2].text = row.get('dataset', '')
            row_cells[3].text = row.get('strengths', '')
            row_cells[4].text = row.get('weaknesses', '')
            
    # Save document to memory stream
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
